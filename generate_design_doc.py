from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_design_document():
    doc = Document()
    
    title = doc.add_heading('资讯日报机器人 - 设计文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. 项目概述', level=1)
    doc.add_paragraph(
        '资讯日报机器人是一个自动化资讯抓取与管理系统，专注于宁德时代（CATL）和小米（XIAOMI）两家公司的新闻资讯。'
        '系统通过多个数据源抓取最新资讯，进行去重、翻译和存储，提供 Web 界面供用户浏览和查询。'
    )
    
    doc.add_heading('1.1 核心功能', level=2)
    doc.add_paragraph('• 手动触发全网资讯同步（最近 7 天）', style='List Bullet')
    doc.add_paragraph('• 多数据源抓取：Google News RSS + GDELT 2.1', style='List Bullet')
    doc.add_paragraph('• 智能去重：基于 URL 和内容哈希', style='List Bullet')
    doc.add_paragraph('• 自动翻译：支持将外文标题和摘要翻译为简体中文', style='List Bullet')
    doc.add_paragraph('• Web 界面：资讯列表浏览、筛选、分页', style='List Bullet')
    doc.add_paragraph('• AI 解读：对资讯进行利多/利空/重要变化分析', style='List Bullet')
    doc.add_paragraph('• 状态监控：实时显示同步状态和统计信息', style='List Bullet')
    
    doc.add_heading('2. 技术架构', level=1)
    
    doc.add_heading('2.1 技术栈', level=2)
    doc.add_paragraph('前端框架：Next.js 16.1.6 (App Router)', style='List Bullet')
    doc.add_paragraph('UI 框架：Tailwind CSS 4', style='List Bullet')
    doc.add_paragraph('数据库：Supabase (PostgreSQL)', style='List Bullet')
    doc.add_paragraph('时间处理：Luxon', style='List Bullet')
    doc.add_paragraph('RSS 解析：rss-parser', style='List Bullet')
    doc.add_paragraph('部署平台：Vercel', style='List Bullet')
    
    doc.add_heading('2.2 系统架构图', level=2)
    doc.add_paragraph('系统采用前后端分离架构，主要包含以下层次：')
    doc.add_paragraph('• 数据层：Supabase PostgreSQL 数据库', style='List Bullet')
    doc.add_paragraph('• 服务层：Next.js API Routes 处理业务逻辑', style='List Bullet')
    doc.add_paragraph('• 数据源层：Google News RSS、GDELT 2.1 API', style='List Bullet')
    doc.add_paragraph('• AI 服务层：智谱 AI / OpenAI 翻译和解读', style='List Bullet')
    doc.add_paragraph('• 展示层：Next.js Server Components + Client Components', style='List Bullet')
    
    doc.add_heading('3. 数据库设计', level=1)
    
    doc.add_heading('3.1 表结构', level=2)
    
    doc.add_heading('3.1.1 job_state 表', level=3)
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    data1 = [
        ('key', 'text (PK)', '任务标识，如 "daily_news"'),
        ('last_success_at', 'timestamptz', '上次成功执行时间'),
        ('updated_at', 'timestamptz', '更新时间'),
    ]
    for row_data in data1:
        row_cells = table1.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
    
    doc.add_heading('3.1.2 run_log 表', level=3)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    data2 = [
        ('id', 'uuid (PK)', '运行记录 ID'),
        ('started_at', 'timestamptz', '开始时间'),
        ('ended_at', 'timestamptz', '结束时间'),
        ('status', 'text', '状态：SUCCESS/FAILED/RUNNING'),
        ('window_start', 'timestamptz', '同步窗口开始时间'),
        ('window_end', 'timestamptz', '同步窗口结束时间'),
        ('fetched_count', 'int', '抓取条数'),
        ('deduped_count', 'int', '去重命中数'),
        ('output_count', 'int', '新增条目数'),
        ('error_message', 'text', '错误信息'),
    ]
    for row_data in data2:
        row_cells = table2.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
    
    doc.add_heading('3.1.3 news_item 表', level=3)
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    data3 = [
        ('id', 'uuid (PK)', '资讯 ID'),
        ('topic', 'text', '主题：CATL/XIAOMI'),
        ('title', 'text', '原始标题'),
        ('title_zh', 'text', '中文标题'),
        ('url', 'text', '文章链接（唯一索引）'),
        ('source', 'text', '来源'),
        ('published_at', 'timestamptz', '发布时间'),
        ('content_hash', 'text', '内容哈希（唯一索引）'),
        ('language', 'text', '语言代码'),
        ('summary', 'text', '原始摘要'),
        ('summary_zh', 'text', '中文摘要'),
        ('created_at', 'timestamptz', '创建时间'),
    ]
    for row_data in data3:
        row_cells = table3.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
    
    doc.add_heading('3.1.4 ai_digest_job 表', level=3)
    table4 = doc.add_table(rows=1, cols=3)
    table4.style = 'Table Grid'
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '说明'
    
    data4 = [
        ('id', 'uuid (PK)', '任务 ID'),
        ('topic', 'text', '主题：CATL/XIAOMI'),
        ('days', 'text', '时间范围：1/7/30/ALL'),
        ('q', 'text', '搜索关键词'),
        ('status', 'text', '状态：PENDING/RUNNING/SUCCESS/FAILED'),
        ('run_token', 'text', '执行令牌'),
        ('result', 'jsonb', '解读结果'),
        ('error_message', 'text', '错误信息'),
        ('created_at', 'timestamptz', '创建时间'),
        ('updated_at', 'timestamptz', '更新时间'),
    ]
    for row_data in data4:
        row_cells = table4.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
    
    doc.add_heading('4. 功能模块', level=1)
    
    doc.add_heading('4.1 资讯同步模块', level=2)
    doc.add_paragraph('手动触发同步，抓取最近 7 天的资讯。')
    doc.add_paragraph('• 入口：/api/sync (POST)', style='List Bullet')
    doc.add_paragraph('• 认证：CRON_SECRET 口令验证', style='List Bullet')
    doc.add_paragraph('• 数据源：Google News RSS（中英文）+ GDELT 2.1', style='List Bullet')
    doc.add_paragraph('• 去重策略：URL 唯一索引 + content_hash 去重', style='List Bullet')
    doc.add_paragraph('• 翻译：自动翻译非中文内容', style='List Bullet')
    
    doc.add_heading('4.2 资讯列表模块', level=2)
    doc.add_paragraph('提供资讯浏览、筛选和分页功能。')
    doc.add_paragraph('• 入口：/news', style='List Bullet')
    doc.add_paragraph('• 筛选条件：公司（CATL/XIAOMI）、时间范围、关键词', style='List Bullet')
    doc.add_paragraph('• 分页：支持自定义每页条数（25/50/100）', style='List Bullet')
    doc.add_paragraph('• 显示：标题（中英文）、摘要（中英文）、来源、发布时间', style='List Bullet')
    
    doc.add_heading('4.3 AI 解读模块', level=2)
    doc.add_paragraph('对资讯进行智能分析，生成利多/利空/重要变化解读。')
    doc.add_paragraph('• 入口：/api/ai/digest (POST)', style='List Bullet')
    doc.add_paragraph('• 两阶段筛选：先筛选重要资讯，再进行深度解读', style='List Bullet')
    doc.add_paragraph('• 执行方式：通过 run_token 安全执行', style='List Bullet')
    doc.add_paragraph('• 结果展示：在资讯列表页上方显示解读结果', style='List Bullet')
    
    doc.add_heading('4.4 状态监控模块', level=2)
    doc.add_paragraph('实时显示系统运行状态和统计信息。')
    doc.add_paragraph('• 入口：/', style='List Bullet')
    doc.add_paragraph('• 显示内容：最近运行状态、同步窗口、统计信息、新增列表', style='List Bullet')
    
    doc.add_heading('5. API 接口', level=1)
    
    doc.add_heading('5.1 POST /api/sync', level=2)
    doc.add_paragraph('手动触发资讯同步')
    doc.add_paragraph('• 请求体：{ secret?: string }', style='List Bullet')
    doc.add_paragraph('• 响应：{ ok: boolean, result: SyncResult }', style='List Bullet')
    doc.add_paragraph('• 认证：CRON_SECRET', style='List Bullet')
    
    doc.add_heading('5.2 POST /api/ai/digest', level=2)
    doc.add_paragraph('创建 AI 解读任务')
    doc.add_paragraph('• 请求体：{ topic: "CATL"|"XIAOMI", days: "1"|"7"|"30"|"ALL", q?: string }', style='List Bullet')
    doc.add_paragraph('• 响应：{ id: string, runToken: string }', style='List Bullet')
    
    doc.add_heading('5.3 POST /api/ai/digest/jobs/[id]/run', level=2)
    doc.add_paragraph('执行 AI 解读任务')
    doc.add_paragraph('• 请求体：{ runToken: string }', style='List Bullet')
    doc.add_paragraph('• 响应：{ ok: boolean, result: DigestResult }', style='List Bullet')
    
    doc.add_heading('5.4 GET /api/ai/digest/jobs/[id]', level=2)
    doc.add_paragraph('查询 AI 解读任务状态')
    doc.add_paragraph('• 响应：{ id, topic, days, q, status, result, error_message, created_at, updated_at }', style='List Bullet')
    
    doc.add_heading('6. 部署方案', level=1)
    
    doc.add_heading('6.1 部署平台', level=2)
    doc.add_paragraph('• 前端 + API：Vercel', style='List Bullet')
    doc.add_paragraph('• 数据库：Supabase', style='List Bullet')
    
    doc.add_heading('6.2 环境变量配置', level=2)
    table5 = doc.add_table(rows=1, cols=3)
    table5.style = 'Table Grid'
    hdr_cells = table5.rows[0].cells
    hdr_cells[0].text = '变量名'
    hdr_cells[1].text = '是否必需'
    hdr_cells[2].text = '说明'
    
    env_vars = [
        ('SUPABASE_URL', '必需', 'Supabase 项目 URL'),
        ('SUPABASE_SERVICE_ROLE_KEY', '必需', 'Supabase 服务角色密钥'),
        ('CRON_SECRET', '可选', '同步口令，用于保护 /api/sync 接口'),
        ('TRANSLATE_TO_ZH', '可选', '是否启用翻译，默认 1'),
        ('TRANSLATION_PROVIDER', '可选', '翻译提供商：zhipu/openai'),
        ('ZHIPU_API_KEY', '可选', '智谱 AI API 密钥'),
        ('OPENAI_API_KEY', '可选', 'OpenAI API 密钥'),
        ('AI_DIGEST', '可选', '是否启用 AI 解读，默认 1'),
        ('AI_PROVIDER', '可选', 'AI 解读提供商：zhipu/openai'),
    ]
    for row_data in env_vars:
        row_cells = table5.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
    
    doc.add_heading('6.3 数据库迁移', level=2)
    doc.add_paragraph('按顺序执行以下迁移文件：')
    doc.add_paragraph('• supabase/migrations/001_init.sql', style='List Bullet')
    doc.add_paragraph('• supabase/migrations/002_news_item_translation.sql', style='List Bullet')
    doc.add_paragraph('• supabase/migrations/003_ai_digest_job.sql', style='List Bullet')
    doc.add_paragraph('• supabase/migrations/004_ai_digest_job_run_token.sql', style='List Bullet')
    doc.add_paragraph('• supabase/migrations/005_remove_email_to.sql', style='List Bullet')
    
    doc.add_heading('7. 安全设计', level=1)
    
    doc.add_heading('7.1 认证与授权', level=2)
    doc.add_paragraph('• /api/sync 接口需要 CRON_SECRET 验证', style='List Bullet')
    doc.add_paragraph('• AI 解读任务使用 run_token 安全执行', style='List Bullet')
    doc.add_paragraph('• Supabase 使用 Row Level Security (RLS)', style='List Bullet')
    
    doc.add_heading('7.2 数据安全', level=2)
    doc.add_paragraph('• 敏感信息通过环境变量配置', style='List Bullet')
    doc.add_paragraph('• 数据库连接使用 Service Role Key', style='List Bullet')
    doc.add_paragraph('• API 密钥不暴露在前端代码中', style='List Bullet')
    
    doc.add_heading('8. 性能优化', level=1)
    
    doc.add_heading('8.1 数据库优化', level=2)
    doc.add_paragraph('• news_item 表建立索引：published_at、url、content_hash', style='List Bullet')
    doc.add_paragraph('• run_log 表建立索引：started_at', style='List Bullet')
    doc.add_paragraph('• 使用 estimated count 提升查询性能', style='List Bullet')
    
    doc.add_heading('8.2 API 优化', level=2)
    doc.add_paragraph('• 使用 Promise.allSettled 并行请求多个数据源', style='List Bullet')
    doc.add_paragraph('• 翻译采用批量处理（每批 10 条）', style='List Bullet')
    doc.add_paragraph('• 设置合理的超时时间（20 秒）', style='List Bullet')
    
    doc.add_heading('9. 错误处理', level=1)
    
    doc.add_heading('9.1 同步错误处理', level=2)
    doc.add_paragraph('• 无口令：显示"同步失败（没有输入口令）"', style='List Bullet')
    doc.add_paragraph('• 口令错误：显示"同步失败（口令错误）"', style='List Bullet')
    doc.add_paragraph('• 同步失败：显示具体错误原因', style='List Bullet')
    doc.add_paragraph('• 翻译失败：不影响主流程，采用 best-effort 策略', style='List Bullet')
    
    doc.add_heading('9.2 AI 解读错误处理', level=2)
    doc.add_paragraph('• 任务状态实时更新', style='List Bullet')
    doc.add_paragraph('• 错误信息记录到 error_message 字段', style='List Bullet')
    doc.add_paragraph('• 前端轮询任务状态，自动重试', style='List Bullet')
    
    doc.add_heading('10. 未来扩展', level=1)
    
    doc.add_heading('10.1 功能扩展', level=2)
    doc.add_paragraph('• 支持更多公司主题', style='List Bullet')
    doc.add_paragraph('• 添加更多数据源', style='List Bullet')
    doc.add_paragraph('• 支持用户订阅和推送通知', style='List Bullet')
    doc.add_paragraph('• 添加数据可视化图表', style='List Bullet')
    
    doc.add_heading('10.2 技术优化', level=2)
    doc.add_paragraph('• 引入缓存机制（Redis）', style='List Bullet')
    doc.add_paragraph('• 优化 AI 解读性能', style='List Bullet')
    doc.add_paragraph('• 添加单元测试和集成测试', style='List Bullet')
    doc.add_paragraph('• 实现日志收集和监控', style='List Bullet')
    
    doc.add_page_break()
    doc.add_heading('附录：项目文件结构', level=1)
    doc.add_paragraph('src/')
    doc.add_paragraph('├── app/')
    doc.add_paragraph('│   ├── api/')
    doc.add_paragraph('│   │   ├── sync/route.ts          # 手动同步接口')
    doc.add_paragraph('│   │   ├── ai/digest/            # AI 解读接口')
    doc.add_paragraph('│   │   └── health/route.ts       # 健康检查')
    doc.add_paragraph('│   ├── page.tsx                  # 状态页')
    doc.add_paragraph('│   ├── news/')
    doc.add_paragraph('│   │   ├── page.tsx              # 资讯列表页')
    doc.add_paragraph('│   │   ├── SyncPanel.tsx         # 同步面板')
    doc.add_paragraph('│   │   └── AiDigestPanel.tsx      # AI 解读面板')
    doc.add_paragraph('│   └── layout.tsx')
    doc.add_paragraph('├── server/')
    doc.add_paragraph('│   ├── manualSync.ts             # 手动同步逻辑')
    doc.add_paragraph('│   ├── googleNewsRss.ts          # Google News 抓取')
    doc.add_paragraph('│   ├── gdelt.ts                  # GDELT 抓取')
    doc.add_paragraph('│   ├── translate.ts              # 翻译服务')
    doc.add_paragraph('│   ├── aiDigest.ts               # AI 解读逻辑')
    doc.add_paragraph('│   └── aiDigestJob.ts           # AI 解读任务管理')
    doc.add_paragraph('├── lib/')
    doc.add_paragraph('│   ├── types.ts                  # 类型定义')
    doc.add_paragraph('│   ├── supabaseAdmin.ts          # Supabase 客户端')
    doc.add_paragraph('│   ├── env.ts                    # 环境变量')
    doc.add_paragraph('│   ├── time.ts                   # 时间工具')
    doc.add_paragraph('│   └── hash.ts                   # 哈希工具')
    doc.add_paragraph('└── supabase/migrations/          # 数据库迁移文件')
    
    return doc

if __name__ == '__main__':
    doc = create_design_document()
    doc.save('设计文档.docx')
    print('设计文档已生成：设计文档.docx')
